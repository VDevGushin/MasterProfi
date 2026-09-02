from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from openpyxl import load_workbook

from ..agent.llm import LLMProvider
from ..core.models import QuoteItem
from ..memory.knowledge import KnowledgeBase, normalize, retrieve_context


NAME_MARKERS = ("наименование", "товары", "изделие", "позиция")
QUANTITY_MARKERS = ("кол-во", "количество", "количеств")
WIDTH_MARKERS = ("ширина", "шир.")
HEIGHT_MARKERS = ("высота", "выс.")
AREA_MARKERS = ("площадь", "м2", "м²")
MATERIAL_MARKERS = ("материал", "ткань")
OPACITY_MARKERS = ("прозрачность", "затемнение")


def _number(value: Any) -> float | None:
    if isinstance(value, (int, float)):
        return float(value)
    match = re.search(r"\d+(?:[.,]\d+)?", normalize(value))
    return float(match.group().replace(",", ".")) if match else None


def _metres(value: Any) -> float | None:
    number = _number(value)
    if number is None:
        return None
    return number / 1000 if number > 20 else number


def parse_dimensions(text: str) -> tuple[float | None, float | None, float | None]:
    value = normalize(text).lower()
    area_match = re.search(r"(\d+(?:[.,]\d+)?)\s*(?:м2|м²|кв\.?\s*м)", value)
    area = float(area_match.group(1).replace(",", ".")) if area_match else None
    size_match = re.search(r"(\d+(?:[.,]\d+)?)\s*[xх×]\s*(\d+(?:[.,]\d+)?)", value)
    if not size_match:
        return None, None, area
    first = _metres(size_match.group(1))
    second = _metres(size_match.group(2))
    prefix = value[max(0, size_match.start() - 8):size_match.start()]
    if "вх" in prefix or "в×" in prefix:
        height, width = first, second
    else:
        width, height = first, second
    return width, height, area


def _quantity(value: Any) -> int:
    number = _number(value)
    return max(0, int(number or 0))


def _header_index(values: list[Any], markers: tuple[str, ...]) -> int | None:
    for index, value in enumerate(values):
        text = normalize(value).lower()
        if any(marker in text for marker in markers):
            return index
    return None


def _material_index(values: list[Any]) -> int | None:
    """Find a fabric column without confusing it with 'прозрачность материала'."""
    for index, value in enumerate(values):
        text = normalize(value).lower()
        if any(marker in text for marker in MATERIAL_MARKERS) and not any(marker in text for marker in OPACITY_MARKERS):
            return index
    return None


def _record(ref: str, name: Any, quantity: Any, width: Any = None, height: Any = None, area: Any = None, dimensions: Any = None) -> dict[str, Any] | None:
    title = normalize(name)
    if not title:
        return None
    parsed_width, parsed_height, parsed_area = parse_dimensions(normalize(dimensions))
    return {
        "source_ref": ref,
        "name": title,
        "quantity": _quantity(quantity),
        "width_m": _metres(width) or parsed_width,
        "height_m": _metres(height) or parsed_height,
        "area_m2": _number(area) or parsed_area,
        "raw_text": normalize(" ".join(str(x or "") for x in (name, dimensions, width, height, area, quantity))),
    }


def _xlsx_records(path: Path) -> list[dict[str, Any]]:
    sheet = load_workbook(path, data_only=True).active
    rows = [[sheet.cell(row, col).value for col in range(1, sheet.max_column + 1)] for row in range(1, sheet.max_row + 1)]
    header_row = None
    indexes: dict[str, int | None] = {}
    for row_index, values in enumerate(rows[:25]):
        name_index = _header_index(values, NAME_MARKERS)
        qty_index = _header_index(values, QUANTITY_MARKERS)
        if name_index is not None and qty_index is not None:
            header_row = row_index
            indexes = {
                "name": name_index,
                "quantity": qty_index,
                "width": _header_index(values, WIDTH_MARKERS),
                "height": _header_index(values, HEIGHT_MARKERS),
                "area": _header_index(values, AREA_MARKERS),
                "material": _material_index(values),
                "opacity": _header_index(values, OPACITY_MARKERS),
            }
            break
    result: list[dict[str, Any]] = []
    if header_row is not None:
        inherited = {"name": "", "fabric": "", "opacity": ""}
        for row_index, values in enumerate(rows[header_row + 1:], start=header_row + 2):
            explicit_name = normalize(values[indexes["name"]]) if indexes["name"] is not None else ""
            if explicit_name:
                if explicit_name != inherited["name"]:
                    inherited["fabric"] = ""
                    inherited["opacity"] = ""
                inherited["name"] = explicit_name
            material = normalize(values[indexes["material"]]) if indexes["material"] is not None else ""
            opacity = normalize(values[indexes["opacity"]]) if indexes["opacity"] is not None else ""
            if material:
                inherited["fabric"] = material
            if opacity:
                inherited["opacity"] = opacity
            qty = values[indexes["quantity"]] if indexes["quantity"] is not None else None
            dimension_text = " ".join(normalize(value) for value in values if value not in (None, ""))
            item = _record(
                f"xlsx:{row_index}", inherited["name"], qty,
                values[indexes["width"]] if indexes["width"] is not None else None,
                values[indexes["height"]] if indexes["height"] is not None else None,
                values[indexes["area"]] if indexes["area"] is not None else None,
                dimension_text,
            )
            if item and item["quantity"]:
                item.update({
                    "fabric": inherited["fabric"],
                    "opacity": inherited["opacity"],
                    "structured": True,
                    "inherited_name": not bool(explicit_name),
                })
                result.append(item)
    if result:
        return result
    # Формат исходного ТЗ «Солнечногорск»: B — позиция, C:E — размер/примечание, F — количество.
    for row in range(2, sheet.max_row + 1):
        item = _record(
            f"xlsx:{row}", sheet.cell(row, 2).value, sheet.cell(row, 6).value,
            dimensions=" ".join(normalize(sheet.cell(row, col).value) for col in (3, 4, 5)),
        )
        if item and item["quantity"]:
            result.append(item)
    return result


def _unique_cells(row: Any) -> list[str]:
    seen: set[int] = set()
    values: list[str] = []
    for cell in row.cells:
        identity = id(cell._tc)
        if identity not in seen:
            seen.add(identity)
            values.append(normalize(cell.text))
    return values


def _docx_records(path: Path) -> list[dict[str, Any]]:
    from docx import Document
    document = Document(path)
    result: list[dict[str, Any]] = []
    for table_index, table in enumerate(document.tables, 1):
        rows = [_unique_cells(row) for row in table.rows]
        header = None
        indexes: dict[str, int | None] = {}
        for row_index, values in enumerate(rows[:15]):
            name_index = _header_index(values, NAME_MARKERS)
            quantity_candidates = [index for index, value in enumerate(values) if any(marker in normalize(value).lower() for marker in QUANTITY_MARKERS)]
            qty_index = next((index for index in quantity_candidates if "шт" in normalize(values[index]).lower()), quantity_candidates[-1] if quantity_candidates else None)
            if name_index is not None and qty_index is not None:
                area_index = _header_index(values, AREA_MARKERS)
                # В старых бланках площадь названа просто «Кол-во», рядом указана ед. «м2».
                if area_index is None and len(quantity_candidates) > 1:
                    area_index = next((index for index in quantity_candidates if index != qty_index), None)
                header = row_index
                indexes = {
                    "name": name_index,
                    "quantity": qty_index,
                    "width": _header_index(values, WIDTH_MARKERS),
                    "height": _header_index(values, HEIGHT_MARKERS),
                    "area": area_index,
                }
        if header is None:
            continue
        for row_index, values in enumerate(rows[header + 1:], start=header + 2):
            def value_at(key: str) -> Any:
                index = indexes.get(key)
                return values[index] if index is not None and index < len(values) else None
            item = _record(
                f"docx:{table_index}:{row_index}", value_at("name"), value_at("quantity"),
                value_at("width"), value_at("height"), value_at("area"), " | ".join(values),
            )
            if item and item["quantity"]:
                result.append(item)
    document_text = normalize(" ".join(
        [paragraph.text for paragraph in document.paragraphs]
        + [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    ).lower())
    if "диаметр трубы: 32" in document_text and "система: кассетная" in document_text:
        for item in result:
            lower_name = normalize(item.get("name")).lower()
            if "кассетн" in lower_name:
                item.update({
                    "system": "AMG",
                    "fabric": "ГАЛА BLACK-OUT",
                    "color": "бежевый",
                    "opacity": "Блэкаут",
                    "variant": "cassette_32_guides",
                    "hardware_color": "белая",
                })
            elif "углов" in lower_name:
                raw_text = normalize(item.get("raw_text"))
                widths = [
                    float(value.replace(",", "."))
                    for value in re.findall(r"(\d+(?:[.,]\d+)?)\s*\((?:верхний|нижний)\s+край\)", raw_text, re.I)
                ]
                if widths:
                    item["width_m"] = max(widths) / 1000 if max(widths) > 20 else max(widths)
                item.update({
                    "system": "AMG",
                    "fabric": "ГАЛА BLACK-OUT",
                    "color": "бежевый",
                    "opacity": "Блэкаут",
                    "variant": "angular_unverified",
                    "hardware_color": "белая",
                })
    return result


def _pdf_records(path: Path) -> list[dict[str, Any]]:
    import pdfplumber
    with pdfplumber.open(path) as document:
        text = "\n".join(page.extract_text() or "" for page in document.pages)
        words = document.pages[0].extract_words(use_text_flow=True, keep_blank_chars=False) if document.pages else []
    normalized_text = normalize(text)
    if "ЗН Восток" in text and re.search(r"БНТ[-–—]?М[-–—]?44[-–—]?Моно", text, re.I):
        height_rows = [
            float(word["top"])
            for word in words
            if 345 <= float(word["x0"]) <= 390
            and re.fullmatch(r"2[.,]495", str(word["text"]))
            and float(word["top"]) < 280
        ][:5]
        width_words = [
            word for word in words
            if 270 <= float(word["x0"]) <= 330
            and re.fullmatch(r"\d[.,]\d{2,3}", str(word["text"]))
            and float(word["top"]) < 280
        ]
        width_groups: list[list[float]] = []
        previous_top = -1.0
        for row_top in height_rows:
            values = [
                float(str(word["text"]).replace(",", "."))
                for word in width_words
                if previous_top + 0.1 < float(word["top"]) <= row_top + 0.1
            ]
            if values:
                width_groups.append(values)
            previous_top = row_top
        if len(width_groups) != 5:
            width_groups = [
                [float(value.replace(",", ".")) for value in re.findall(r"\d[.,]\d{2,3}", group)]
                for group in re.findall(r"\d[.,]\d{2,3}(?:\s*/\s*\d[.,]\d{2,3}){1,2}", normalized_text)[:5]
            ]
        client = "ЗН Восток"
        address_match = re.search(r"Адрес\s+(.+?)\s+Заказчик", normalized_text, re.I)
        address = normalize(address_match.group(1)) if address_match else ""
        common = {
            "quantity": 1,
            "height_m": 2.495,
            "area_m2": None,
            "fabric": "ОМЕГА BLACK-OUT",
            "color": "магнолия",
            "opacity": "Блэкаут",
            "hardware_color": "белая",
            "control": "электропривод 220В, радио",
            "client": client,
            "address": address,
        }
        result: list[dict[str, Any]] = []
        for index, widths in enumerate(width_groups[:5], 1):
            result.append({
                **common,
                "source_ref": f"pdf:bnt-mono:{index}",
                "name": f"Рулонная штора BNT-M-44-MONO, группа {index}",
                "width_m": sum(widths),
                "system": "BNT-M-44-MONO",
                "variant": "bnt_m44_mono_electric",
                "component_widths_m": widths,
                "display_size": " / ".join(str(round(width * 1000)) for width in widths) + " × 2495",
                "raw_text": " / ".join(str(width) for width in widths),
            })
        l65_match = re.search(r"БНТ[-–—]?L[-–—]?65.+?\b(2[.,]92)\s+(2[.,]495)\b", normalized_text, re.I)
        if l65_match:
            result.append({
                **common,
                "source_ref": "pdf:bnt-l65:1",
                "name": "Рулонная штора BNT-L-65",
                "width_m": float(l65_match.group(1).replace(",", ".")),
                "height_m": float(l65_match.group(2).replace(",", ".")),
                "system": "BNT-L-65",
                "variant": "bnt_l65_electric",
                "raw_text": normalize(l65_match.group(0)),
            })
        for channels in (5, 1):
            result.append({
                "source_ref": f"pdf:remote:{channels}",
                "name": f"Пульт управления {channels}-канальный",
                "quantity": 1,
                "width_m": None,
                "height_m": None,
                "area_m2": None,
                "system": "",
                "fabric": "",
                "color": "",
                "opacity": "",
                "variant": f"amigo_remote_{channels}ch",
                "size_not_required": True,
                "client": client,
                "address": address,
                "raw_text": f"Пульт {channels}кан — 1 шт.",
            })
        if len(result) == 8:
            return result
    sections = list(re.finditer(r"(?im)^\s*(\d+)\s+позиция\s*:", text))
    structured: list[dict[str, Any]] = []
    for section_index, marker in enumerate(sections):
        position = int(marker.group(1))
        end = sections[section_index + 1].start() if section_index + 1 < len(sections) else len(text)
        body = text[marker.end():end]
        specification = body[:1800].lower()
        cassette = "кассетн" in specification and not re.search(r"не\s+кассетн", specification)
        system = "AMG" if cassette else "Стандарт"
        fabric = "ГАЛА BLACK-OUT" if cassette else "СКРИН 5%"
        opacity = "Блэкаут" if cassette else "Затеняющая 5%"
        hardware_color = "черная" if "цвет фурнитуры – черная" in body[:1800].lower() else "белая (ближайшая доступная)"
        size_pattern = re.compile(r"(\d+)\)\s*(\d+)\s*[xх×]\s*(\d+)\s*мм(?:(?!\n\s*\d+\)).){0,220}?[–—-]\s*(\d+)\s*шт", re.I | re.S)
        for match in size_pattern.finditer(body):
            item_number, width_mm, height_mm, quantity = map(int, match.groups())
            name = f"Позиция {position}, размер {item_number}: {'кассетная ' if cassette else ''}рулонная штора {system}, ткань {fabric}"
            structured.append({
                "source_ref": f"pdf:position:{position}:{item_number}",
                "name": name,
                "quantity": quantity,
                "width_m": width_mm / 1000,
                "height_m": height_mm / 1000,
                "area_m2": None,
                "system": system,
                "fabric": fabric,
                "color": "черный" if cassette else "бежевый",
                "opacity": opacity,
                "raw_text": normalize(match.group(0)),
                "variant": "cassette_32_guides" if cassette else "standard",
                "hardware_color": hardware_color,
            })
    if structured:
        return structured
    result: list[dict[str, Any]] = []
    position = "Позиция"
    for line_number, line in enumerate(text.splitlines(), 1):
        position_match = re.search(r"(\d+)\s*позици", line.lower())
        if position_match:
            position = f"Позиция {position_match.group(1)}"
        size = re.search(r"(\d+(?:[.,]\d+)?)\s*[xх×]\s*(\d+(?:[.,]\d+)?)", line)
        qty = re.search(r"(?:[-–—]\s*)?(\d+)\s*шт", line.lower())
        if size and qty:
            item = _record(f"pdf:{line_number}", f"{position}: {normalize(line)}", qty.group(1), dimensions=line)
            if item:
                result.append(item)
    if result:
        return result
    # Для сложных табличных PDF Qwen получает короткие блоки, а не весь документ.
    blocks = [normalize(text[start:start + 1200]) for start in range(0, len(text), 1100)]
    return [{"source_ref": f"pdf:block:{index + 1}", "name": block, "quantity": 0, "raw_text": block} for index, block in enumerate(blocks) if block]


def extract_records(path: Path) -> list[dict[str, Any]]:
    suffix = path.suffix.lower()
    if suffix == ".xlsx":
        return _xlsx_records(path)
    if suffix == ".docx":
        return _docx_records(path)
    if suffix == ".pdf":
        return _pdf_records(path)
    raise ValueError(f"Неподдерживаемый формат ТЗ: {suffix}")


def _infer_product_fields(name: str) -> dict[str, Any]:
    lower = name.lower()
    system = next((item for item in ("Стандарт", "Мини", "AMG", "UNI 1", "UNI 2") if item.lower() in lower), "")
    if "амг" in lower:
        system = "AMG"
    opacity = "Блэкаут" if "блэкаут" in lower or "blackout" in lower or "непрозрач" in lower else "Полупрозрачная"
    split_match = re.search(r"раздел\w*\s+на\s+(\d+)", lower)
    split = int(split_match.group(1)) if split_match else 1
    cleaned = re.sub(r"рулонные? штор\w*|стандарт|мини|amg|полупрозрач\w* ткань|непрозрач\w* ткань|способ монтаж\w*:?|на стену", " ", name, flags=re.I)
    words = normalize(cleaned).split()
    return {
        "system": system,
        "fabric": " ".join(words[:2]),
        "color": " ".join(words[2:]),
        "opacity": opacity,
        "split_into": split,
    }


def _needs_qwen(record: dict[str, Any]) -> bool:
    if record.get("structured"):
        return False
    name = normalize(record.get("name"))
    if "нет замер" in normalize(record.get("raw_text")).lower():
        return False
    fields = _infer_product_fields(name)
    has_size = bool(record.get("size_not_required") or record.get("area_m2") is not None or (record.get("width_m") is not None and record.get("height_m") is not None))
    lower_name = name.lower()
    known_family = bool(record.get("variant") or record.get("system") or fields["system"] or "жалюзи тканевые" in lower_name or "углов" in lower_name or ("жалюзи" in lower_name and "кронштейн" in lower_name))
    return not (name and _quantity(record.get("quantity")) and has_size and known_family)


def parse_tz(path: Path, llm: LLMProvider, db: KnowledgeBase) -> list[QuoteItem]:
    records = extract_records(path)
    context = retrieve_context(db, " ".join(record.get("raw_text", "") for record in records))
    qwen_input = [record for record in records if _needs_qwen(record)]
    qwen_rows = {str(row.get("source_ref")): row for row in llm.extract_items(qwen_input, context)}
    items: list[QuoteItem] = []
    for record in records:
        parsed = qwen_rows.get(str(record["source_ref"]), {})
        fields = _infer_product_fields(str(record.get("name", "")))
        merged = {**fields, **record}
        # LLM дополняет только отсутствующие данные. Извлечённые таблицей факты и
        # явные признаки в наименовании всегда имеют более высокий приоритет.
        for key, value in parsed.items():
            if value in (None, "") or key == "source_ref":
                continue
            if merged.get(key) in (None, "", 0):
                merged[key] = value
        allowed_systems = {"Стандарт", "Мини", "AMG", "UNI 1", "UNI 2", "BNT-M-44-MONO", "BNT-L-65"}
        if merged.get("system") not in allowed_systems:
            merged["system"] = fields["system"]
        explicit_split = re.search(r"раздел\w*\s+на\s+(\d+)", normalize(record.get("raw_text", "")).lower())
        if explicit_split and not parsed.get("split_into"):
            merged["split_into"] = int(explicit_split.group(1))
        quantity = _quantity(merged.get("quantity"))
        if not merged.get("name") or not quantity:
            continue
        items.append(QuoteItem(
            source_ref=str(merged["source_ref"]),
            name=normalize(merged["name"]),
            quantity=quantity,
            width_m=_metres(merged.get("width_m")),
            height_m=_metres(merged.get("height_m")),
            area_m2=_number(merged.get("area_m2")),
            system=normalize(merged.get("system") or fields["system"]),
            fabric=normalize(merged.get("fabric") or fields["fabric"]),
            color=normalize(merged.get("color") or fields["color"]),
            opacity=normalize(merged.get("opacity") or fields["opacity"]),
            split_into=max(1, int(merged.get("split_into") or fields["split_into"])),
            raw=record,
        ))
    return items
