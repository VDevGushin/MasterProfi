from __future__ import annotations

import hashlib
import json
import re
import sqlite3
try:
    import tomllib
except ModuleNotFoundError:  # Python 3.9–3.10
    import tomli as tomllib
from datetime import datetime
from pathlib import Path
from typing import Any

from openpyxl import load_workbook

from .config import COMPETITORS_CONFIG, COMPETITOR_SOURCES, EXAMPLES_DIR, KNOWLEDGE_DB, KNOWLEDGE_DIR, REFERENCE_QUOTES_DIR


def normalize(value: Any) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip()


def normalize_key(value: Any) -> str:
    return normalize(value).upper().replace("Ё", "Е")


class KnowledgeBase:
    def __init__(self) -> None:
        KNOWLEDGE_DIR.mkdir(parents=True, exist_ok=True)
        self.connection = sqlite3.connect(KNOWLEDGE_DB)
        self.connection.execute("""
            CREATE TABLE IF NOT EXISTS knowledge (
                kind TEXT NOT NULL,
                key TEXT NOT NULL,
                payload TEXT NOT NULL,
                source TEXT NOT NULL,
                confidence REAL NOT NULL,
                verified INTEGER NOT NULL,
                updated_at TEXT NOT NULL,
                PRIMARY KEY(kind, key)
            )
        """)
        self.connection.commit()

    def put(self, kind: str, key: str, payload: dict[str, Any], source: str, confidence: float = 1.0, verified: bool = False) -> None:
        self.connection.execute("""
            INSERT INTO knowledge(kind,key,payload,source,confidence,verified,updated_at)
            VALUES(?,?,?,?,?,?,?)
            ON CONFLICT(kind,key) DO UPDATE SET payload=excluded.payload, source=excluded.source,
              confidence=excluded.confidence, verified=excluded.verified, updated_at=excluded.updated_at
        """, (kind, normalize_key(key), json.dumps(payload, ensure_ascii=False), source, confidence, int(verified), datetime.now().isoformat()))
        self.connection.commit()

    def search(self, kind: str, query: str, limit: int = 5, verified_only: bool = False) -> list[dict[str, Any]]:
        sql = "SELECT key,payload,source,confidence,verified FROM knowledge WHERE kind=?"
        params: list[Any] = [kind]
        if verified_only:
            sql += " AND verified=1"
        query_tokens = set(normalize_key(query).split())
        ranked: list[tuple[int, dict[str, Any]]] = []
        for key, payload, source, confidence, verified in self.connection.execute(sql, params):
            data = json.loads(payload)
            haystack = set((key + " " + normalize_key(data.get("text", "")) + " " + normalize_key(data.get("chunk", ""))).split())
            score = len(query_tokens & haystack)
            if normalize_key(query) == key:
                score += 100
            if score:
                ranked.append((score, {"key": key, "payload": data, "source": source, "confidence": confidence, "verified": bool(verified)}))
        ranked.sort(key=lambda item: item[0], reverse=True)
        return [item for _, item in ranked[:limit]]

    def close(self) -> None:
        self.connection.close()


def load_verified_analogues(db: KnowledgeBase) -> int:
    count = 0
    if not COMPETITORS_CONFIG.exists():
        return count
    with COMPETITORS_CONFIG.open("rb") as stream:
        analogues = tomllib.load(stream).get("analogue", [])
    for item in analogues:
        if not item.get("verified") or not item.get("source_material"):
            continue
        db.put("analogue", str(item["source_material"]), item, "config/competitors.toml", float(item.get("confidence", 1.0)), True)
        count += 1
    return count


def load_competitor_sources(db: KnowledgeBase) -> int:
    if not COMPETITOR_SOURCES.exists():
        return 0
    with COMPETITOR_SOURCES.open("rb") as stream:
        sources = tomllib.load(stream).get("source", [])
    for item in sources:
        db.put("competitor_source", item["url"], {**item, "text": " ".join(str(v) for v in item.values())}, "source/knowledge/competitor_sources.toml", 1.0, True)
    return len(sources)


def _extract_reference(path: Path) -> str:
    if path.suffix.lower() == ".xlsx":
        book = load_workbook(path, data_only=True, read_only=True)
        lines: list[str] = []
        for sheet in book.worksheets:
            for row in sheet.iter_rows():
                values = [normalize(cell.value) for cell in row if cell.value not in (None, "")]
                if values:
                    lines.append(" | ".join(values))
        return "\n".join(lines)
    if path.suffix.lower() == ".docx":
        from docx import Document
        doc = Document(path)
        lines = [normalize(p.text) for p in doc.paragraphs if normalize(p.text)]
        for table in doc.tables:
            for row in table.rows:
                seen: set[int] = set()
                values: list[str] = []
                for cell in row.cells:
                    identity = id(cell._tc)
                    if identity not in seen and normalize(cell.text):
                        seen.add(identity)
                        values.append(normalize(cell.text))
                if values:
                    lines.append(" | ".join(values))
        return "\n".join(lines)
    if path.suffix.lower() == ".pdf":
        import pdfplumber
        with pdfplumber.open(path) as document:
            return "\n".join(page.extract_text() or "" for page in document.pages)
    return ""


def refresh_rag(db: KnowledgeBase) -> dict[str, int]:
    files = 0
    chunks = 0
    for directory, label in ((EXAMPLES_DIR, "example_tz"), (REFERENCE_QUOTES_DIR, "example_quote")):
        if not directory.exists():
            continue
        for path in sorted(directory.iterdir()):
            if path.suffix.lower() not in {".xlsx", ".docx", ".pdf"}:
                continue
            try:
                text = normalize(_extract_reference(path))
            except Exception:
                continue
            if not text:
                continue
            files += 1
            digest = hashlib.sha1(path.read_bytes()).hexdigest()[:10]
            for index, start in enumerate(range(0, len(text), 900)):
                chunk = text[start:start + 1100]
                db.put("example", f"{label}:{path.name}:{digest}:{index}", {"file": path.name, "kind": label, "chunk": chunk}, str(path), 1.0, True)
                chunks += 1
    return {"files": files, "chunks": chunks}


def retrieve_context(db: KnowledgeBase, query: str, limit_chars: int = 1400) -> str:
    fragments: list[str] = []
    seen: set[str] = set()
    for entry in db.search("example", query, limit=12, verified_only=True):
        data = entry["payload"]
        filename = str(data.get("file", ""))
        if filename in seen:
            continue
        seen.add(filename)
        fragments.append(f"Пример {filename}: {normalize(data.get('chunk', ''))[:420]}")
        if len(fragments) == 3:
            break
    return "\n".join(fragments)[:limit_chars]


def initialize_knowledge(db: KnowledgeBase) -> dict[str, Any]:
    return {
        "verified_analogues": load_verified_analogues(db),
        "competitor_sources": load_competitor_sources(db),
        "rag": refresh_rag(db),
    }
